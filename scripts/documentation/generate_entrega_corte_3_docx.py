"""
Generador de documento Word (APA) para Entrega Corte 3 - Multicont
Incluye: Wireframes (con imágenes y descripciones), Paleta de Colores, Árbol de Navegación,
Público objetivo, Entidad, Cantidad de usuarios, Características de usuarios y Accesibilidad.

Salida: docs/business/wireframes/ENTREGA_CORTE_3_APA.docx
Fecha: 2025-10-28
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = "c:/Users/spiri/MultiContGit"
WF_DIR = os.path.join(ROOT, "docs", "business", "wireframes")
OUTPUT = os.path.join(WF_DIR, "ENTREGA_CORTE_3_APA.docx")

# Imagenes en orden
WIREFRAMES = [
    ("WF-001_login.png", "Figura 1. Wireframe de Login", "Pantalla de inicio de sesión con campos de usuario, contraseña, recordar sesión y enlace de recuperación."),
    ("WF-002_dashboard.png", "Figura 2. Dashboard Principal", "Panel con KPIs, gráfica de ventas, top productos y alertas, con navegación lateral."),
    ("WF-003_organizations_list.png", "Figura 3. Lista de Organizaciones", "Vista tipo tabla con búsqueda, filtros, paginación y acciones CRUD."),
    ("WF-004_organization_form.png", "Figura 4. Formulario de Organización (Modal)", "Modal para crear/editar organización, con validaciones y estado."),
    ("WF-005_inventory_list.png", "Figura 5. Lista de Inventario", "Control de stock con alertas visuales de bajo inventario y acciones por fila."),
    ("WF-006_create_quote.png", "Figura 6. Crear Cotización", "Formulario multi-línea con cálculo de totales e interacción dinámica de líneas."),
    ("WF-007_analytics_dashboard.png", "Figura 7. Dashboard de Analytics", "Métricas, filtros, gráficos comparativos y rankings (exportar Excel/PDF)."),
    ("WF-008_users_management.png", "Figura 8. Gestión de Usuarios", "CRUD de usuarios con roles y estados, filtros y paginación (solo ADMIN)."),
]

# Contenidos clave (extraídos de WIREFRAMES_NUEVOS.md)
PALETA = [
    ("Azul Corporativo", "#1E40AF", "Headers, botones primarios"),
    ("Gris Oscuro", "#1F2937", "Texto principal, sidebar"),
    ("Blanco", "#FFFFFF", "Fondos, cards"),
    ("Azul Claro", "#3B82F6", "Hover, links"),
    ("Verde Éxito", "#10B981", "Éxito/confirmaciones"),
    ("Amarillo Alerta", "#F59E0B", "Advertencias/pending"),
    ("Rojo Peligro", "#EF4444", "Errores/eliminar"),
    ("Gris Neutro", "#6B7280", "Texto secundario/disabled"),
]

PUBLICO = (
    "Empresas medianas y grandes con múltiples sucursales (comercio, distribución, servicios industriales). "
    "Tamaños por empresa: 10–20 (pequeña), 50–100 (mediana), 200–500 (grande). "
    "Total multi-tenant estimado: hasta 10,000 usuarios concurrentes. Distribución de roles (cada 100): "
    "ADMIN 2–3%, MANAGER 10–15%, SALES 60–70%, VIEWER 15–20%."
)

CARACTERISTICAS_USUARIOS = [
    "Conocimiento tecnológico: ADMIN (alto), MANAGER (medio-alto), SALES (básico-medio), VIEWER (medio).",
    "Dispositivos de acceso: Desktop 70%, Laptop 20%, Tablet 8%, Móvil 2%; navegadores modernos (Chrome/Edge/Firefox/Safari).",
    "Modo responsivo: enfoque mobile-first con breakpoints sm, md, lg, xl, 2xl; sidebar colapsable en <1024px.",
    "Conectividad: optimizado para 3–10 Mbps con paginación, compresión (gzip), cache y lazy loading de gráficos.",
]

ACCESIBILIDAD = [
    "Objetivo WCAG 2.1 Nivel AA.",
    "Contraste verificado (ej.: #1F2937 sobre #FFFFFF = 16.1:1).",
    "Navegación por teclado (Tab, Shift+Tab, Enter, Escape, flechas).",
    "Etiquetas ARIA y roles semánticos en componentes interactivos.",
    "Compatibilidad con lectores de pantalla (NVDA, JAWS, VoiceOver, TalkBack).",
    "Tamaños mínimos: fuente base 14px, botones 44×44px, zoom 200% sin pérdida.",
    "Alternativas al color (iconografía, patrones, texto descriptivo).",
]

ARBOL_FILE = os.path.join(WF_DIR, "ARBOL_NAVEGACION.md")

def set_apa_defaults(document: Document):
    # Márgenes 1" y fuente Times New Roman 12, doble espaciado
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    # Fuente normal
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


def add_title_page(document: Document):
    # Portada APA simple (estudiante)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Entrega Corte 3 — Diseño Visual y Navegación\n")
    run.bold = True
    run.font.size = Pt(20)
    p.add_run("Sistema de Gestión Empresarial Multicont\n")
    p.add_run("Wilker & Daniel\n")
    p.add_run("Curso: Desarrollo de Aplicaciones Web\n")
    p.add_run("Fecha: 28 de Octubre de 2025")
    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_bullets(document: Document, items):
    for it in items:
        p = document.add_paragraph(it, style='List Bullet')
        p_format = p.paragraph_format
        p_format.line_spacing = 2


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph(text)
    p.paragraph_format.line_spacing = 2
    return p


def add_palette_table(document: Document):
    add_paragraph(document, "Paleta de colores corporativa (resumen):")
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Color'
    hdr_cells[1].text = 'Hex'
    hdr_cells[2].text = 'Uso'
    for name, hex_code, use in PALETA:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = hex_code
        row_cells[2].text = use
    add_paragraph(document, "")


def add_wireframe_image(document: Document, path: str, caption: str, desc: str, fig_num: int):
    if os.path.exists(path):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(6.5))
    # Caption APA (centrada, en cursiva)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    # Descripción
    add_paragraph(document, desc)


def extract_arbol_codeblock(md_path: str) -> str:
    if not os.path.exists(md_path):
        return ""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    # Buscar primer bloque de código triple ``` y extraer
    start = content.find("```")
    if start == -1:
        return content
    end = content.find("```", start + 3)
    if end == -1:
        return content
    return content[start+3:end].strip('\n')


def add_code_block(document: Document, text: str):
    # Simular bloque monoespaciado
    for line in text.splitlines():
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        p.paragraph_format.line_spacing = 1.15


def build_document():
    doc = Document()
    set_apa_defaults(doc)
    add_title_page(doc)

    # 1. Wireframes + paleta + disposición
    add_heading(doc, "Wireframes de la aplicación y paleta de colores", level=1)
    add_paragraph(doc, (
        "Los wireframes siguen una estructura con header fijo (logo, búsqueda, usuario, notificaciones, ajustes), "
        "sidebar de navegación por módulos y área central para contenido (cards, tablas, formularios y gráficos)."
    ))
    add_palette_table(doc)

    # Insertar imágenes en orden con captions
    add_heading(doc, "Wireframes (imágenes y descripción)", level=2)
    for idx, (fname, caption, desc) in enumerate(WIREFRAMES, start=1):
        img_path = os.path.join(WF_DIR, fname)
        add_heading(doc, caption.replace("Figura", "WF"), level=3)
        add_wireframe_image(doc, img_path, caption, desc, idx)

    doc.add_page_break()

    # 2. Árbol de navegación
    add_heading(doc, "Árbol de navegación de la aplicación", level=1)
    add_paragraph(doc, "A continuación se incluye el árbol de navegación completo, que cubre las áreas funcionales, administrativas y puntos de acceso a vistas:")
    tree_text = extract_arbol_codeblock(ARBOL_FILE)
    add_code_block(doc, tree_text)

    doc.add_page_break()

    # 3. Público objetivo, entidad y cantidad de usuarios
    add_heading(doc, "Público objetivo, entidad y cantidad de usuarios", level=1)
    add_paragraph(doc, PUBLICO)

    # 4. Características de los usuarios
    add_heading(doc, "Características de los usuarios", level=1)
    add_bullets(doc, CARACTERISTICAS_USUARIOS)

    # 5. Accesibilidad
    add_heading(doc, "Características de accesibilidad consideradas", level=1)
    add_bullets(doc, ACCESIBILIDAD)

    # Guardar
    doc.save(OUTPUT)
    print(f"✅ Documento generado: {OUTPUT}")


if __name__ == '__main__':
    build_document()
