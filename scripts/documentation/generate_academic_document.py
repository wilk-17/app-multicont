"""
Script para generar documento Word académico
Formato APA 7 con contenido de:
- ARBOL_DE_PROBLEMAS.md (Planteamiento)
- ARBOL_DE_OBJETIVOS.md (Justificación + Objetivos)
- METODOLOGIA_RAD.md (Metodología)

Ejecutar: python generate_academic_document.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import os


def create_apa_styles(doc):
    """Crear estilos personalizados según APA 7"""
    
    # Estilo para títulos de nivel 1 (centrado, bold, 14pt)
    try:
        styles = doc.styles
        
        # Title style
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(14)
        title_font.bold = True
        
        # Heading 1 (centrado, bold, 12pt)
        h1_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(12)
        h1_font.bold = True
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading 2 (izquierda, bold, 12pt)
        h2_style = styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Normal text (12pt, double-spaced)
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 2.0  # Double spacing
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    except Exception as e:
        print(f"Warning: Could not apply all styles - {e}")


def add_cover_page(doc):
    """Agregar portada según APA 7"""
    
    # Título (centrado, mitad superior)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SISTEMA DE GESTIÓN EMPRESARIAL MULTICONT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Planteamiento del Problema, Justificación, Objetivos y Metodología RAD")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Espaciado
    for _ in range(3):
        doc.add_paragraph()
    
    # Autores (centrado, parte inferior)
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(
        "Daniel Romero\n"
        "Raquel Morales\n"
        "David Piñeros\n"
        "William Wilches"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Institución
    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institution.add_run("Universidad Nacional de Colombia\nFacultad de Ingeniería")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Fecha
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"{date.today().strftime('%d de %B de %Y')}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()


def add_planteamiento(doc):
    """Agregar sección de Planteamiento del Problema"""
    
    # Título de sección (Heading 1)
    heading = doc.add_heading('PLANTEAMIENTO DEL PROBLEMA', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introducción
    p = doc.add_paragraph(
        "Las organizaciones empresariales con múltiples sucursales enfrentan desafíos "
        "significativos en la gestión coordinada de sus operaciones. La falta de sistemas "
        "integrados genera ineficiencias operativas, pérdida de información crítica, y "
        "dificultades en el seguimiento del desempeño de empleados y sucursales."
    )
    
    # Árbol de Problemas
    doc.add_heading('Árbol de Problemas Identificados', level=2)
    
    p = doc.add_paragraph(
        "El análisis del contexto empresarial reveló cuatro problemas fundamentales que "
        "impactan negativamente la operación de organizaciones multilocales:"
    )
    
    # Problema 1
    doc.add_heading('Problema 1: Pérdida de Control de Inventario', level=2)
    p = doc.add_paragraph(
        "La ausencia de trazabilidad en las asignaciones de items de inventario a empleados "
        "genera pérdidas económicas estimadas en el 15% del valor total del inventario anual. "
        "Los empleados reciben equipos (laptops, herramientas, muestras de producto) sin un "
        "registro formal que permita auditar su ubicación, estado y responsable actual."
    )
    p = doc.add_paragraph(
        "Impacto cuantificado: Para una organización con inventario valorado en $500,000,000 COP, "
        "esto representa pérdidas anuales de $75,000,000 COP."
    )
    
    # Problema 2
    doc.add_heading('Problema 2: Dificultad para Medir Desempeño de Ventas', level=2)
    p = doc.add_paragraph(
        "La carencia de un sistema de metas de ventas con seguimiento automatizado impide "
        "evaluar objetivamente el rendimiento de empleados y sucursales. Los gerentes dependen "
        "de hojas de cálculo Excel desactualizadas, lo que genera cálculos erróneos del "
        "porcentaje de cumplimiento y retrasos en la toma de decisiones."
    )
    p = doc.add_paragraph(
        "Impacto cuantificado: El tiempo invertido en consolidación manual de datos representa "
        "40 horas/mes por sucursal, equivalente a $3,200,000 COP anuales en costos administrativos "
        "para una red de 5 sucursales."
    )
    
    # Problema 3
    doc.add_heading('Problema 3: Análisis Ineficiente de Facturación', level=2)
    p = doc.add_paragraph(
        "Sin herramientas de analytics, las organizaciones no pueden identificar patrones de "
        "venta por empleado, sucursal o marca de producto. Esta falta de visibilidad impide "
        "optimizar el inventario, reasignar recursos, y diseñar estrategias comerciales basadas "
        "en datos."
    )
    p = doc.add_paragraph(
        "Impacto cuantificado: Las decisiones comerciales basadas en intuición en lugar de datos "
        "generan una pérdida de oportunidades estimada en $12,500,000 COP anuales por no identificar "
        "productos de alta rotación."
    )
    
    # Problema 4
    doc.add_heading('Problema 4: Información Fragmentada en Múltiples Sistemas', level=2)
    p = doc.add_paragraph(
        "Las organizaciones operan con sistemas desconectados: un software para facturación, "
        "hojas de Excel para inventario, emails para cotizaciones. Esta fragmentación causa "
        "errores de transcripción, duplicación de datos, y baja productividad del personal."
    )
    p = doc.add_paragraph(
        "Impacto cuantificado: Los errores de facturación por datos inconsistentes generan "
        "reprocesos que cuestan aproximadamente $8,000,000 COP anuales en tiempo de corrección "
        "y pérdida de confianza del cliente."
    )
    
    doc.add_page_break()


def add_justificacion(doc):
    """Agregar sección de Justificación"""
    
    heading = doc.add_heading('JUSTIFICACIÓN', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introducción
    p = doc.add_paragraph(
        "El desarrollo del Sistema Multicont se justifica por la necesidad imperativa de "
        "modernizar las operaciones empresariales mediante tecnología que centralice, automatice "
        "y optimice procesos críticos de negocio."
    )
    
    # Justificación Económica
    doc.add_heading('Justificación Económica', level=2)
    p = doc.add_paragraph(
        "El análisis costo-beneficio demuestra que la implementación del sistema genera un "
        "Retorno de Inversión (ROI) altamente positivo:"
    )
    
    # Tabla de costos
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Monto Anual (COP)'
    
    # Data
    data = [
        ('Reducción pérdidas inventario (50% de $75M)', '$37,500,000'),
        ('Ahorro en tiempo administrativo', '$3,200,000'),
        ('Mejora en decisiones comerciales', '$12,500,000'),
        ('Reducción de errores de facturación', '$8,000,000'),
        ('TOTAL AHORRO ANUAL', '$90,500,000')
    ]
    
    for i, (concept, amount) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = concept
        row_cells[1].text = amount
    
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        "Con un costo de desarrollo estimado en $30,000,000 COP y costos operativos anuales "
        "de $5,000,000 COP (hosting + mantenimiento), el sistema se amortiza en menos de 5 meses "
        "y genera un ahorro neto de $55,500,000 COP en el primer año."
    )
    
    # Justificación Tecnológica
    doc.add_heading('Justificación Tecnológica', level=2)
    p = doc.add_paragraph(
        "La arquitectura del sistema se fundamenta en principios de Clean Architecture (Robert C. Martin, 2017), "
        "que garantiza mantenibilidad, escalabilidad y testability del código. La separación en tres capas "
        "(Entities, Use Cases, API) permite modificar componentes sin afectar la lógica de negocio central."
    )
    
    p = doc.add_paragraph(
        "La elección de Flask (framework Python) y PostgreSQL (base de datos relacional) se basa en: "
        "(a) madurez tecnológica con comunidad activa, (b) escalabilidad probada en sistemas empresariales, "
        "(c) licencias open-source que reducen costos de licenciamiento, y (d) facilidad de integración con "
        "herramientas de Business Intelligence."
    )
    
    # Justificación Metodológica
    doc.add_heading('Justificación Metodológica', level=2)
    p = doc.add_paragraph(
        "La adopción de la metodología RAD (Rapid Application Development) se justifica por la necesidad "
        "de entregar funcionalidad de negocio rápidamente en ciclos iterativos. A diferencia del modelo "
        "en cascada tradicional, RAD permite feedback temprano del cliente, reducción de riesgos por "
        "validación incremental, y flexibilidad para ajustar requerimientos."
    )
    
    p = doc.add_paragraph(
        "Como evidencia, el proyecto se desarrolló en 6 fases iterativas de 3 semanas cada una, con "
        "entregables funcionales al final de cada fase. Esta aproximación resultó en 80+ endpoints REST "
        "completamente documentados y testeados en un período de 5 meses."
    )
    
    doc.add_page_break()


def add_objetivos(doc):
    """Agregar sección de Objetivos"""
    
    heading = doc.add_heading('OBJETIVOS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Objetivo General
    doc.add_heading('Objetivo General', level=2)
    p = doc.add_paragraph(
        "Desarrollar un sistema de gestión empresarial integral basado en arquitectura limpia "
        "que permita a organizaciones multilocales centralizar, automatizar y optimizar procesos "
        "de inventario, ventas, facturación y recursos humanos, mejorando la eficiencia operativa "
        "en al menos un 40% y reduciendo costos administrativos en $90,500,000 COP anuales."
    )
    
    # Objetivos Específicos
    doc.add_heading('Objetivos Específicos', level=2)
    
    objectives = [
        {
            'num': 1,
            'title': 'Implementar módulo de trazabilidad de asignaciones',
            'desc': 'Desarrollar sistema de tracking de items de inventario asignados a empleados '
                    'con estados (activo, devuelto, perdido) y registro de condición, reduciendo '
                    'pérdidas por falta de control en un 50%.',
            'metric': 'Métrica de éxito: 100% de asignaciones rastreables con historial auditable.'
        },
        {
            'num': 2,
            'title': 'Desarrollar sistema de metas de ventas con períodos configurables',
            'desc': 'Crear módulo que permita asignar metas mensuales, trimestrales y anuales a '
                    'empleados o sucursales, con comparación automática contra facturación real y '
                    'cálculo de porcentaje de cumplimiento.',
            'metric': 'Métrica de éxito: 18+ metas configuradas con estados dinámicos (exceeded, '
                     'on_track, at_risk, failed).'
        },
        {
            'num': 3,
            'title': 'Implementar analytics avanzados de facturación',
            'desc': 'Desarrollar 7 endpoints especializados de analytics que permitan analizar '
                    'facturación por empleado, sucursal y marca, identificando top performers y '
                    'tendencias de venta.',
            'metric': 'Métrica de éxito: Reducción del 80% en tiempo de generación de reportes '
                     '(de 40 horas/mes a 8 horas/mes).'
        },
        {
            'num': 4,
            'title': 'Diseñar arquitectura escalable con Clean Architecture',
            'desc': 'Aplicar principios de Clean Architecture con separación en 3 capas (Entities, '
                    'Use Cases, API) para garantizar mantenibilidad y escalabilidad del código.',
            'metric': 'Métrica de éxito: 85%+ de cobertura de tests unitarios y cero dependencias '
                     'circulares entre capas.'
        },
        {
            'num': 5,
            'title': 'Implementar seguridad con RBAC y JWT',
            'desc': 'Desarrollar sistema de autenticación basado en tokens JWT con Control de Acceso '
                    'por Roles (4 niveles: ADMIN, MANAGER, SALES, VIEWER) para proteger información '
                    'sensible.',
            'metric': 'Métrica de éxito: 100% de endpoints protegidos con decoradores @require_role().'
        },
        {
            'num': 6,
            'title': 'Documentar APIs con estándar Swagger/OpenAPI',
            'desc': 'Generar documentación interactiva de 80+ endpoints REST usando Flasgger, '
                    'facilitando integración con sistemas frontend.',
            'metric': 'Métrica de éxito: 100% de endpoints documentados con ejemplos de request/response.'
        },
        {
            'num': 7,
            'title': 'Validar sistema mediante testing exhaustivo',
            'desc': 'Desarrollar 45+ test cases con pytest cubriendo CRUD completo, integración de '
                    'módulos y casos edge, garantizando calidad del código.',
            'metric': 'Métrica de éxito: 85%+ de cobertura de tests sin errores críticos en producción.'
        }
    ]
    
    for obj in objectives:
        doc.add_heading(f"Objetivo Específico {obj['num']}", level=3)
        p = doc.add_paragraph(obj['desc'])
        p = doc.add_paragraph(f"✓ {obj['metric']}", style='List Bullet')
    
    doc.add_page_break()


def add_metodologia(doc):
    """Agregar sección de Metodología"""
    
    heading = doc.add_heading('METODOLOGÍA', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introducción
    p = doc.add_paragraph(
        "El proyecto Sistema Multicont se desarrolló aplicando la metodología RAD (Rapid Application Development), "
        "una aproximación iterativa e incremental que prioriza la entrega rápida de funcionalidad de negocio "
        "mediante ciclos cortos de desarrollo (timeboxes) con feedback continuo del usuario final."
    )
    
    # Fundamentos de RAD
    doc.add_heading('Fundamentos de la Metodología RAD', level=2)
    p = doc.add_paragraph(
        "RAD, propuesta por James Martin en 1991, se caracteriza por cuatro principios clave:"
    )
    
    principles = [
        "Desarrollo iterativo: El sistema se construye en incrementos funcionales, no en una sola entrega monolítica.",
        "Prototipos funcionales: Cada iteración entrega código ejecutable que puede ser probado por stakeholders.",
        "Timeboxing: Ciclos de desarrollo de duración fija (3 semanas en este proyecto) que fuerzan priorización.",
        "Participación activa del usuario: Feedback temprano permite ajustar requerimientos sin retrasos costosos."
    ]
    
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    # Comparación con Cascada
    doc.add_heading('Diferencias con el Modelo en Cascada', level=2)
    
    # Tabla comparativa
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aspecto'
    hdr_cells[1].text = 'Cascada (Waterfall)'
    hdr_cells[2].text = 'RAD (Aplicado)'
    
    # Data
    data = [
        ('Entregas', 'Una sola entrega final', '6 entregas incrementales'),
        ('Duración de fases', 'Fases largas (meses)', 'Fases cortas (3 semanas)'),
        ('Flexibilidad', 'Cambios costosos', 'Cambios bienvenidos'),
        ('Testing', 'Al final del proyecto', 'En cada iteración')
    ]
    
    for i, (aspect, waterfall, rad) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = aspect
        row_cells[1].text = waterfall
        row_cells[2].text = rad
    
    doc.add_paragraph()
    
    # Fases Aplicadas
    doc.add_heading('Fases del Proyecto Multicont', level=2)
    
    phases = [
        {
            'num': 1,
            'name': 'Fundamentos y Autenticación',
            'duration': '3 semanas',
            'deliverables': 'Entidades User, Role, Permission. JWT authentication. 5 tests unitarios.'
        },
        {
            'num': 2,
            'name': 'Gestión Organizacional',
            'duration': '3 semanas',
            'deliverables': 'Entidades Organization, Branch, Employee, Person. CRUD completo. 12 tests.'
        },
        {
            'num': 3,
            'name': 'Inventario y Asignaciones',
            'duration': '3 semanas',
            'deliverables': 'Entidades InventoryItem, Brand, Category, Assignment. Trazabilidad implementada. 10 tests.'
        },
        {
            'num': 4,
            'name': 'Flujo de Ventas',
            'duration': '3 semanas',
            'deliverables': 'Entidades Quote, SalesOrder, Invoice con items. Conversión Quote→Order→Invoice. 15 tests.'
        },
        {
            'num': 5,
            'name': 'Sistema de Metas',
            'duration': '3 semanas',
            'deliverables': 'Entidad SalesGoal. Períodos configurables. Comparación metas vs actual. 8 tests.'
        },
        {
            'num': 6,
            'name': 'Analytics Avanzados',
            'duration': '3 semanas',
            'deliverables': '7 endpoints de analytics. Cache implementado. Documentación Swagger completa. 10 tests.'
        }
    ]
    
    for phase in phases:
        doc.add_heading(f"Fase {phase['num']}: {phase['name']}", level=3)
        p = doc.add_paragraph(f"Duración: {phase['duration']}")
        p = doc.add_paragraph(f"Entregables: {phase['deliverables']}")
    
    # Evidencia de Aplicación
    doc.add_heading('Evidencia de Aplicación de RAD', level=2)
    p = doc.add_paragraph(
        "La aplicación exitosa de RAD en el proyecto Multicont se evidencia en:"
    )
    
    evidence = [
        "80+ endpoints REST funcionales desarrollados en 5 meses (16 endpoints/mes).",
        "6 entregas incrementales con código ejecutable y documentado.",
        "45+ test cases que validan funcionalidad en cada iteración.",
        "Commits atómicos en GitHub con mensajes descriptivos (100+ commits).",
        "Documentación técnica actualizada en cada fase (25+ documentos MD).",
        "Diagramas UML y wireframes generados iterativamente (14 archivos PNG)."
    ]
    
    for item in evidence:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph(
        "Esta evidencia demuestra inequívocamente que el proyecto NO siguió un modelo en cascada, "
        "sino una aproximación iterativa con entregas frecuentes y validación continua."
    )
    
    doc.add_page_break()


def add_referencias(doc):
    """Agregar sección de Referencias (APA 7)"""
    
    heading = doc.add_heading('REFERENCIAS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Referencias en formato APA 7
    references = [
        "Martin, R. C. (2017). Clean architecture: A craftsman's guide to software structure and design. Prentice Hall.",
        
        "Martin, J. (1991). Rapid application development. Macmillan Publishing Co., Inc.",
        
        "Cockburn, A. (2005). Hexagonal architecture. Alistair Cockburn's website. https://alistair.cockburn.us/hexagonal-architecture/",
        
        "Fowler, M. (2002). Patterns of enterprise application architecture. Addison-Wesley.",
        
        "Evans, E. (2003). Domain-driven design: Tackling complexity in the heart of software. Addison-Wesley Professional.",
        
        "Ronacher, A. (2010). Flask: A Python microframework. Pallets Projects. https://flask.palletsprojects.com/",
        
        "PostgreSQL Global Development Group. (2023). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/index.html"
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
        p.paragraph_format.line_spacing = 2.0


def generate_document():
    """Función principal para generar el documento"""
    
    print("📄 Generando documento académico...")
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes (1 pulgada en todos lados - APA 7)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Aplicar estilos APA
    create_apa_styles(doc)
    
    # Agregar contenido
    print("  ✓ Generando portada...")
    add_cover_page(doc)
    
    print("  ✓ Agregando Planteamiento del Problema...")
    add_planteamiento(doc)
    
    print("  ✓ Agregando Justificación...")
    add_justificacion(doc)
    
    print("  ✓ Agregando Objetivos...")
    add_objetivos(doc)
    
    print("  ✓ Agregando Metodología RAD...")
    add_metodologia(doc)
    
    print("  ✓ Agregando Referencias (APA 7)...")
    add_referencias(doc)
    
    # Guardar documento
    output_path = os.path.join('docs', 'PLANTEAMIENTO_PROYECTO.docx')
    os.makedirs('docs', exist_ok=True)
    doc.save(output_path)
    
    print(f"\n✅ Documento generado exitosamente: {output_path}")
    print(f"   Páginas aproximadas: 12-15")
    print(f"   Formato: APA 7")
    print(f"   Contenido: Planteamiento, Justificación, Objetivos, Metodología, Referencias")


if __name__ == '__main__':
    try:
        generate_document()
    except ImportError:
        print("\n❌ ERROR: Falta el paquete python-docx")
        print("Instalar con: pip install python-docx")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
